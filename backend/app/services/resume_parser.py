"""
Resume ingestion service.

Supports PDF (pdfplumber) and DOCX (python-docx).
Extracts structured fields using regex + heuristics, then wraps
everything in a ParsedResume schema.
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path

from app.schemas.resume import ParsedResume

logger = logging.getLogger(__name__)

# ── Skill vocabularies ────────────────────────────────────────────────────────

TECH_SKILLS = {
    # Languages
    "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
    "Kotlin", "Swift", "PHP", "Ruby", "Scala", "R", "MATLAB", "Bash", "Shell",
    # Web
    "React", "Vue", "Angular", "Next.js", "Nuxt.js", "Svelte", "HTML", "CSS",
    "SASS", "Tailwind", "Bootstrap", "jQuery",
    # Backend
    "FastAPI", "Django", "Flask", "Express", "Spring Boot", "Node.js", "Rails",
    ".NET", "ASP.NET", "Laravel",
    # Data / ML
    "PyTorch", "TensorFlow", "Keras", "scikit-learn", "Pandas", "NumPy",
    "Spark", "Hadoop", "Airflow", "dbt", "MLflow",
    # Databases
    "PostgreSQL", "MySQL", "SQLite", "MongoDB", "Redis", "Elasticsearch",
    "Cassandra", "DynamoDB", "BigQuery", "Snowflake",
    # Cloud / DevOps
    "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "Ansible",
    "Jenkins", "GitHub Actions", "CircleCI", "Prometheus", "Grafana",
    # Messaging
    "Kafka", "RabbitMQ", "SQS", "Pub/Sub",
    # APIs
    "REST", "GraphQL", "gRPC", "OpenAPI",
    # Other
    "Git", "Linux", "Agile", "Scrum", "CI/CD",
}

SOFT_SKILLS = {
    "communication", "leadership", "teamwork", "collaboration", "problem-solving",
    "critical thinking", "time management", "adaptability", "creativity",
    "project management", "mentoring", "conflict resolution",
}

SENIORITY_KEYWORDS = {
    "executive": ["cto", "ceo", "vp ", "vice president", "director", "head of"],
    "lead":      ["lead", "staff", "principal", "architect"],
    "senior":    ["senior", "sr.", "sr "],
    "mid":       ["mid-level", "intermediate"],
    "junior":    ["junior", "jr.", "jr ", "entry-level", "entry level"],
    "internship":["intern", "co-op", "coop"],
}


def _infer_seniority(text: str) -> str:
    tl = text.lower()
    for level, keywords in SENIORITY_KEYWORDS.items():
        if any(k in tl for k in keywords):
            return level
    return "mid"


def _extract_experience_years(text: str) -> float | None:
    """Look for patterns like '5+ years', '3 years of experience'."""
    patterns = [
        r"(\d+)\+?\s+years?\s+of\s+(?:professional\s+)?experience",
        r"(\d+)\+?\s+years?\s+(?:of\s+)?(?:work|industry|relevant)",
        r"experience[:\s]+(\d+)\+?\s+years?",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return float(m.group(1))
    return None


def _extract_skills(text: str) -> tuple[list[str], list[str], list[str]]:
    """Returns (tech_skills, soft_skills, all_skills)."""
    text_lower = text.lower()
    tech: list[str] = []
    soft: list[str] = []

    for skill in TECH_SKILLS:
        pattern = re.escape(skill)
        if re.search(r"\b" + pattern + r"\b", text, re.IGNORECASE):
            tech.append(skill)

    for skill in SOFT_SKILLS:
        if skill.lower() in text_lower:
            soft.append(skill)

    all_skills = list(dict.fromkeys(tech + soft))  # preserve order, dedupe
    return tech, soft, all_skills


def _extract_education(text: str) -> list[str]:
    degrees = []
    patterns = [
        r"(Bachelor(?:'s)?(?:\s+of\s+\w+)?(?:\s+in\s+[^\n,]+)?)",
        r"(Master(?:'s)?(?:\s+of\s+\w+)?(?:\s+in\s+[^\n,]+)?)",
        r"(PhD|Ph\.D\.?|Doctorate)(?:\s+in\s+[^\n,]+)?",
        r"(Diploma\s+in\s+[^\n,]+)",
        r"(Certificate\s+in\s+[^\n,]+)",
        r"(B\.?Sc\.?|B\.?Eng\.?|B\.?A\.?|M\.?Sc\.?|M\.?Eng\.?|MBA)(?:\s+[^\n,]+)?",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            edu = m.group(0).strip()
            if edu not in degrees:
                degrees.append(edu)
    return degrees[:5]


def _extract_certifications(text: str) -> list[str]:
    cert_patterns = [
        r"(AWS\s+Certified\s+[^\n,]+)",
        r"(Azure\s+[^\n,]+Certification[^\n,]*)",
        r"(GCP\s+[^\n,]+)",
        r"(Google\s+Cloud\s+[^\n,]+)",
        r"(PMP|CISSP|CEH|CPA|CFA|CISA|CISM|CompTIA\s+\w+)",
        r"(Certified\s+[^\n,]+)",
        r"(Kubernetes\s+(?:CKAD|CKA|CKS)[^\n,]*)",
    ]
    certs: list[str] = []
    for pat in cert_patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            cert = m.group(0).strip()
            if cert not in certs:
                certs.append(cert)
    return certs[:10]


def _extract_job_titles(text: str) -> list[str]:
    """Grab the most prominent job titles from experience sections."""
    title_patterns = [
        r"(?:^|\n)([A-Z][a-zA-Z\s]+(?:Engineer|Developer|Designer|Manager|"
        r"Analyst|Architect|Lead|Director|Scientist|Consultant|Specialist|"
        r"Officer|Administrator|Coordinator))[,\|\n]",
    ]
    titles: list[str] = []
    for pat in title_patterns:
        for m in re.finditer(pat, text, re.MULTILINE):
            t = m.group(1).strip()
            if t not in titles and len(t) < 80:
                titles.append(t)
    return titles[:6]


# ── File parsers ─────────────────────────────────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_text_from_txt(file_bytes: bytes) -> str:
    import chardet

    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding") or "utf-8"
    return file_bytes.decode(encoding, errors="replace")


# ── Public API ────────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    """
    Parse a resume file (PDF/DOCX/TXT) and return a ParsedResume.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        raw_text = _extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        raw_text = _extract_text_from_docx(file_bytes)
    elif ext in (".txt", ".md"):
        raw_text = _extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported resume format: {ext}")

    if not raw_text.strip():
        raise ValueError("Could not extract text from resume — file may be image-based.")

    tech_skills, soft_skills, all_skills = _extract_skills(raw_text)
    education = _extract_education(raw_text)
    certifications = _extract_certifications(raw_text)
    job_titles = _extract_job_titles(raw_text)
    experience_years = _extract_experience_years(raw_text)
    seniority = _infer_seniority(raw_text)

    return ParsedResume(
        raw_text=raw_text,
        skills=all_skills,
        tech_stack=tech_skills,
        soft_skills=soft_skills,
        experience_years=experience_years,
        seniority_level=seniority,
        job_titles=job_titles,
        education=education,
        certifications=certifications,
    )
